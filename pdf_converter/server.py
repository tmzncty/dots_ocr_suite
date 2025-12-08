#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF to DOCX Converter Server
使用 http.server 提供静态文件和API服务
"""

import http.server
import json
import os
import hashlib
import zipfile
import traceback
import time
from pathlib import Path
from urllib.parse import unquote, quote
from multiprocessing import Pool
import threading
import sys

# 添加父目录到路径以导入库
sys.path.insert(0, str(Path(__file__).parent.parent))

from dots_ocr_lib import DotsOCRParser, load_images_from_pdf

# Markdown to DOCX
from docx import Document
from docx.shared import Pt
import base64
import io
import re

# ==============================================================================
# Configuration
# ==============================================================================
PORT = 7860
BASE_DIR = Path(__file__).parent
STATIC_DIR = BASE_DIR / "static"
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(exist_ok=True)

# 初始化 Parser
parser = DotsOCRParser(
    ip="192.168.24.78",
    port=8000,
    dpi=150,
    min_pixels=3136,
    max_pixels=11289600
)

# 处理状态存储
processing_state = {}

# ==============================================================================
# Helper Functions
# ==============================================================================

def log_to_state(hash_id, message):
    """添加日志到处理状态"""
    if hash_id in processing_state:
        processing_state[hash_id]['log'] = message
        print(f"[{hash_id}] {message}")

def get_file_hash(file_data, length=8):
    """获取文件的SHA256哈希"""
    return hashlib.sha256(file_data).hexdigest()[:length]

def process_single_page(args):
    """处理单个页面（多进程）"""
    origin_image, save_dir, save_name, page_idx, hash_id = args
    
    log_to_state(hash_id, f"正在处理第 {page_idx + 1} 页...")
    
    result = parser._parse_single_image(
        origin_image=origin_image,
        prompt_mode='prompt_layout_all_en',
        save_dir=str(save_dir),
        save_name=save_name,
        source='pdf',
        page_idx=page_idx
    )
    
    return result

def markdown_to_docx(md_content, output_path):
    """将Markdown转换为DOCX"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=min(level, 9))
        
        # 图片
        elif line.startswith('![') and '](data:image/' in line:
            try:
                match = re.search(r'!\[.*?\]\((data:image/[^;]+;base64,[^)]+)\)', line)
                if match:
                    data_url = match.group(1)
                    header, encoded = data_url.split(',', 1)
                    image_data = base64.b64decode(encoded)
                    
                    image_stream = io.BytesIO(image_data)
                    try:
                        from docx.shared import Inches
                        doc.add_picture(image_stream, width=Inches(5))
                    except:
                        doc.add_paragraph('[Image]')
            except Exception as e:
                doc.add_paragraph(f'[Image - Error: {str(e)}]')
        
        # LaTeX公式
        elif line.startswith('$$'):
            formula_lines = [line]
            i += 1
            while i < len(lines) and not lines[i].strip().endswith('$$'):
                formula_lines.append(lines[i])
                i += 1
            if i < len(lines):
                formula_lines.append(lines[i])
            formula_text = '\n'.join(formula_lines)
            p = doc.add_paragraph(formula_text)
            p.style = 'Intense Quote'
        
        # 普通段落
        else:
            text = line
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            doc.add_paragraph(text)
        
        i += 1
    
    doc.save(output_path)

def create_zip_package(base_dir, base_name, hash_id):
    """创建包含DOCX、MD、JSON的ZIP包"""
    zip_name = f"{base_name}_{hash_id}.zip"
    zip_path = base_dir / zip_name
    
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        docx_file = base_dir / f"{base_name}_{hash_id}.docx"
        if docx_file.exists():
            zipf.write(docx_file, f"{base_name}_{hash_id}.docx")
        
        md_file = base_dir / f"{base_name}_{hash_id}_combined.md"
        if md_file.exists():
            zipf.write(md_file, f"{base_name}_{hash_id}.md")
        
        json_file = base_dir / f"{base_name}_{hash_id}_combined.json"
        if json_file.exists():
            zipf.write(json_file, f"{base_name}_{hash_id}.json")
    
    return zip_path

def process_pdf_background(pdf_path, work_dir, base_name, hash_id, process_mode, filename):
    """后台处理PDF"""
    start_time = time.time()
    
    try:
        # 1. 拆图阶段
        log_to_state(hash_id, "开始从PDF提取页面...")
        processing_state[hash_id].update({
            'extract_progress': 10,
            'extract_status': 'Loading PDF...'
        })
        
        images = load_images_from_pdf(str(pdf_path), dpi=150)
        total_pages = len(images)
        
        log_to_state(hash_id, f"PDF加载成功，共 {total_pages} 页")
        processing_state[hash_id].update({
            'extract_progress': 50,
            'extract_status': f'Loaded {total_pages} pages'
        })
        
        # 根据模式决定处理哪些页面
        if process_mode == 'single':
            images_to_process = [images[0]]
            page_indices = [0]
            log_to_state(hash_id, "处理模式: 仅处理第1页")
        else:
            images_to_process = images
            page_indices = list(range(total_pages))
            log_to_state(hash_id, f"处理模式: 处理全部 {total_pages} 页")
        
        processing_state[hash_id].update({
            'extract_progress': 100,
            'extract_status': 'Complete'
        })
        
        # 2. OCR处理阶段
        log_to_state(hash_id, f"开始OCR处理 {len(images_to_process)} 页...")
        processing_state[hash_id].update({
            'ocr_progress': 0,
            'ocr_status': 'Starting OCR...'
        })
        
        # 多进程处理
        args_list = [
            (img, str(work_dir), base_name, idx, hash_id)
            for img, idx in zip(images_to_process, page_indices)
        ]
        
        results = []
        with Pool() as pool:
            for i, result in enumerate(pool.imap(process_single_page, args_list)):
                results.append(result)
                progress = (i + 1) / len(args_list) * 100
                processing_state[hash_id].update({
                    'ocr_progress': progress,
                    'ocr_status': f'Page {i+1}/{len(args_list)}'
                })
                log_to_state(hash_id, f"完成OCR处理: 第 {i+1}/{len(args_list)} 页")
        
        processing_state[hash_id].update({
            'ocr_progress': 100,
            'ocr_status': 'Complete'
        })
        
        # 3. 生成文档阶段
        log_to_state(hash_id, "开始合并结果并生成文档...")
        processing_state[hash_id].update({
            'generate_progress': 10,
            'generate_status': 'Merging results...'
        })
        
        all_cells = []
        all_md_parts = []
        
        for result in sorted(results, key=lambda x: x['page_no']):
            # 读取JSON
            with open(result['layout_info_path'], 'r', encoding='utf-8') as f:
                cells = json.load(f)
                all_cells.append({'page': result['page_no'], 'cells': cells})
            
            # 读取MD
            with open(result['md_content_path'], 'r', encoding='utf-8') as f:
                md_content = f.read()
                all_md_parts.append(f"# Page {result['page_no'] + 1}\n\n{md_content}")
        
        processing_state[hash_id].update({
            'generate_progress': 40,
            'generate_status': 'Saving JSON...'
        })
        
        # 保存合并的JSON
        combined_json_path = work_dir / f"{base_name}_{hash_id}_combined.json"
        with open(combined_json_path, 'w', encoding='utf-8') as f:
            json.dump(all_cells, f, ensure_ascii=False, indent=2)
        
        log_to_state(hash_id, "合并的JSON已保存")
        
        processing_state[hash_id].update({
            'generate_progress': 60,
            'generate_status': 'Saving Markdown...'
        })
        
        # 保存合并的MD
        combined_md = '\n\n---\n\n'.join(all_md_parts)
        combined_md_path = work_dir / f"{base_name}_{hash_id}_combined.md"
        with open(combined_md_path, 'w', encoding='utf-8') as f:
            f.write(combined_md)
        
        log_to_state(hash_id, "合并的Markdown已保存")
        
        processing_state[hash_id].update({
            'generate_progress': 75,
            'generate_status': 'Creating DOCX...'
        })
        
        # 生成DOCX
        docx_path = work_dir / f"{base_name}_{hash_id}.docx"
        markdown_to_docx(combined_md, str(docx_path))
        
        log_to_state(hash_id, "DOCX文档已生成")
        
        processing_state[hash_id].update({
            'generate_progress': 90,
            'generate_status': 'Creating ZIP...'
        })
        
        # 创建ZIP
        create_zip_package(work_dir, base_name, hash_id)
        
        log_to_state(hash_id, "ZIP包已创建")
        
        processing_time = f"{time.time() - start_time:.2f}s"
        
        # 完成
        processing_state[hash_id].update({
            'generate_progress': 100,
            'generate_status': 'Complete',
            'complete': True,
            'filename': filename,
            'total_pages': total_pages,
            'hash_id': hash_id,
            'processing_time': processing_time
        })
        
        log_to_state(hash_id, f"所有处理完成！总耗时: {processing_time}")
    
    except Exception as e:
        error_msg = str(e)
        log_to_state(hash_id, f"处理失败: {error_msg}")
        processing_state[hash_id].update({
            'complete': True,
            'error': error_msg
        })
        print(f"Processing error: {traceback.format_exc()}")

# ==============================================================================
# Server Handler
# ==============================================================================
class PDFConverterHandler(http.server.BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        """自定义日志输出"""
        print(f"{self.address_string()} - {format%args}")
    
    def do_GET(self):
        try:
            # 健康检查
            if self.path == '/ping':
                self.send_response(200)
                self.end_headers()
                self.wfile.write(b'pong')
                return

            # API端点
            if self.path.startswith('/progress/'):
                hash_id = self.path.split('/')[-1]
                state = processing_state.get(hash_id, {
                    'extract_progress': 0,
                    'extract_status': 'Unknown',
                    'ocr_progress': 0,
                    'ocr_status': 'Unknown',
                    'generate_progress': 0,
                    'generate_status': 'Unknown',
                    'complete': False
                })
                
                response_data = json.dumps(state).encode('utf-8')
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.send_header('Content-Length', str(len(response_data)))
                self.end_headers()
                self.wfile.write(response_data)
                return
            
            elif self.path.startswith('/download/'):
                parts = self.path.split('/')
                if len(parts) < 4:
                    self.send_error(400, "Invalid request")
                    return
                    
                hash_id = parts[2]
                file_type = parts[3]
                
                for item in DATA_DIR.iterdir():
                    if item.is_dir() and hash_id in item.name:
                        base_name = item.name.replace(f"_{hash_id}", "")
                        
                        if file_type == 'zip':
                            zip_path = item / f"{base_name}_{hash_id}.zip"
                            if zip_path.exists():
                                self.send_file(zip_path, 'application/zip')
                                return
                        elif file_type == 'docx':
                            docx_path = item / f"{base_name}_{hash_id}.docx"
                            if docx_path.exists():
                                self.send_file(docx_path, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                                return
                
                self.send_error(404, "File not found")
                return
            
            elif self.path == '/list_files':
                files = []
                for item in DATA_DIR.iterdir():
                    if item.is_dir():
                        name_parts = item.name.rsplit('_', 1)
                        if len(name_parts) == 2:
                            base_name, hash_id = name_parts
                            json_file = item / f"{base_name}_{hash_id}_combined.json"
                            pages = 0
                            if json_file.exists():
                                try:
                                    with open(json_file, 'r', encoding='utf-8') as f:
                                        data = json.load(f)
                                        pages = len(data)
                                except:
                                    pass
                            
                            files.append({
                                'name': base_name,
                                'hash_id': hash_id,
                                'pages': pages
                            })
                
                response_data = json.dumps({'files': files}).encode('utf-8')
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.send_header('Content-Length', str(len(response_data)))
                self.end_headers()
                self.wfile.write(response_data)
                return
            
            # 静态文件处理
            path = self.path
            if path == '/':
                path = '/index.html'
            
            # 移除查询参数
            path = path.split('?')[0]
            
            # 安全检查：防止目录遍历
            if '..' in path:
                self.send_error(403, "Forbidden")
                return
            
            # 移除开头的 /
            if path.startswith('/'):
                path = path[1:]
            
            file_path = STATIC_DIR / path
            
            if file_path.exists() and file_path.is_file():
                # 确定 MIME 类型
                ext = file_path.suffix.lower()
                content_type = 'application/octet-stream'
                if ext == '.html':
                    content_type = 'text/html; charset=utf-8'
                elif ext == '.css':
                    content_type = 'text/css'
                elif ext == '.js':
                    content_type = 'application/javascript'
                elif ext == '.json':
                    content_type = 'application/json'
                elif ext == '.png':
                    content_type = 'image/png'
                elif ext == '.jpg' or ext == '.jpeg':
                    content_type = 'image/jpeg'
                
                with open(file_path, 'rb') as f:
                    content = f.read()
                    self.send_response(200)
                    self.send_header('Content-type', content_type)
                    self.send_header('Content-Length', str(len(content)))
                    self.end_headers()
                    self.wfile.write(content)
            else:
                self.send_error(404, f"File not found: {path}")
        
        except Exception as e:
            print(f"Error in do_GET: {traceback.format_exc()}")
            self.send_error(500, f"Internal server error: {str(e)}")
    
    def send_file(self, file_path, content_type):
        """发送文件"""
        try:
            print(f"Sending file: {file_path}")
            if not file_path.exists():
                print(f"File not found: {file_path}")
                self.send_error(404, "File not found")
                return

            with open(file_path, 'rb') as f:
                content = f.read()
                self.send_response(200)
                self.send_header('Content-type', content_type)
                
                # 处理文件名
                filename = file_path.name
                # URL编码文件名
                encoded_name = quote(filename)
                
                # 打印调试信息
                print(f"Original filename: {filename}")
                print(f"Encoded filename: {encoded_name}")
                
                # 设置 Content-Disposition
                # 优先使用 filename*=UTF-8'' 格式
                header_value = f"attachment; filename*=UTF-8''{encoded_name}"
                
                # 尝试添加 ASCII 兼容的 filename 参数（可选，为了兼容旧浏览器）
                # safe_filename = filename.encode('ascii', 'ignore').decode('ascii') or "download.file"
                # header_value += f'; filename="{safe_filename}"'
                
                print(f"Content-Disposition: {header_value}")
                self.send_header('Content-Disposition', header_value)
                
                self.send_header('Content-Length', str(len(content)))
                self.end_headers()
                self.wfile.write(content)
                print("File sent successfully")
        except Exception as e:
            print(f"Error sending file: {traceback.format_exc()}")
            self.send_error(500, f"Error sending file: {str(e)}")
    
    def do_POST(self):
        try:
            if self.path == '/upload_and_process':
                content_type = self.headers['Content-Type']
                if 'multipart/form-data' not in content_type:
                    raise ValueError("Invalid content type")
                
                boundary = content_type.split('boundary=')[1].encode()
                content_length = int(self.headers['Content-Length'])
                post_data = self.rfile.read(content_length)
                
                # 解析multipart数据
                parts = post_data.split(b'--' + boundary)
                file_data = None
                filename = None
                process_mode = 'all'
                
                for part in parts:
                    if b'Content-Disposition' in part:
                        if b'filename=' in part:
                            lines = part.split(b'\r\n')
                            for line in lines:
                                if b'filename=' in line:
                                    filename = line.decode('utf-8', errors='ignore').split('filename="')[1].split('"')[0]
                                    break
                            file_data = part.split(b'\r\n\r\n', 1)[1].rsplit(b'\r\n', 1)[0]
                        elif b'name="process_mode"' in part:
                            process_mode = part.split(b'\r\n\r\n', 1)[1].rsplit(b'\r\n', 1)[0].decode()
                
                if not file_data or not filename:
                    raise ValueError("No file uploaded")
                
                # 计算哈希
                hash_id = get_file_hash(file_data)
                base_name = Path(filename).stem
                
                print(f"Processing: {filename} (hash: {hash_id}, mode: {process_mode})")
                
                # 创建工作目录
                work_dir = DATA_DIR / f"{base_name}_{hash_id}"
                
                # 检查是否已存在
                if work_dir.exists():
                    docx_path = work_dir / f"{base_name}_{hash_id}.docx"
                    if docx_path.exists():
                        json_file = work_dir / f"{base_name}_{hash_id}_combined.json"
                        pages = 0
                        if json_file.exists():
                            with open(json_file, 'r', encoding='utf-8') as f:
                                pages = len(json.load(f))
                        
                        response = {
                            'already_exists': True,
                            'hash_id': hash_id,
                            'filename': filename,
                            'total_pages': pages
                        }
                        
                        response_data = json.dumps(response).encode('utf-8')
                        self.send_response(200)
                        self.send_header('Content-type', 'application/json')
                        self.send_header('Content-Length', str(len(response_data)))
                        self.end_headers()
                        self.wfile.write(response_data)
                        return
                
                work_dir.mkdir(exist_ok=True)
                
                # 保存PDF
                pdf_path = work_dir / filename
                with open(pdf_path, 'wb') as f:
                    f.write(file_data)
                
                # 初始化处理状态
                processing_state[hash_id] = {
                    'extract_progress': 0,
                    'extract_status': 'Starting...',
                    'ocr_progress': 0,
                    'ocr_status': 'Waiting...',
                    'generate_progress': 0,
                    'generate_status': 'Waiting...',
                    'complete': False,
                    'log': 'Initialized'
                }
                
                # 启动后台处理
                thread = threading.Thread(
                    target=process_pdf_background,
                    args=(pdf_path, work_dir, base_name, hash_id, process_mode, filename)
                )
                thread.daemon = True
                thread.start()
                
                response = {'hash_id': hash_id, 'already_exists': False}
                response_data = json.dumps(response).encode('utf-8')
                
                self.send_response(200)
                self.send_header('Content-type', 'application/json')
                self.send_header('Content-Length', str(len(response_data)))
                self.end_headers()
                self.wfile.write(response_data)
                return
            
            self.send_error(404, "Not found")
        
        except Exception as e:
            print(f"Error in do_POST: {traceback.format_exc()}")
            error_response = json.dumps({"error": str(e)}).encode('utf-8')
            self.send_response(500)
            self.send_header('Content-type', 'application/json')
            self.send_header('Content-Length', str(len(error_response)))
            self.end_headers()
            self.wfile.write(error_response)

# ==============================================================================
# Main
# ==============================================================================
def run(port=PORT):
    server_address = ('0.0.0.0', port)
    # 使用 ThreadingHTTPServer 支持并发请求
    if hasattr(http.server, 'ThreadingHTTPServer'):
        server_class = http.server.ThreadingHTTPServer
    else:
        # 兼容旧版本
        import socketserver
        class ThreadingHTTPServer(socketserver.ThreadingMixIn, http.server.HTTPServer):
            pass
        server_class = ThreadingHTTPServer
        
    httpd = server_class(server_address, PDFConverterHandler)
    
    print("=" * 60)
    print(f"PDF to DOCX Converter Server")
    print("=" * 60)
    print(f"Server running on port {port}")
    print(f"Open: http://localhost:{port}")
    print(f"Data directory: {DATA_DIR.resolve()}")
    print(f"Static files: {STATIC_DIR.resolve()}")
    print("=" * 60)
    print("Press Ctrl+C to stop the server")
    print()
    
    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        print("\nShutting down server...")
        httpd.shutdown()

if __name__ == "__main__":
    run()
